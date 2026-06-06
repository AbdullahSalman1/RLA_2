from docx import Document
from docx.shared import Pt

filename = "Personal_Journey.docx"

doc = Document()

def add_heading(text, level=1):
    if level == 1:
        doc.add_heading(text, level=1)
    else:
        doc.add_heading(text, level=2)


def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold

# Title
add_heading("Personal Journey — Project: Distance", level=1)

# Introduction
add_paragraph(
    "Introduction:\n"
    "This document summarizes my personal journey while working on the `Distance` project. It covers the project's goals, the difficulties I faced, how I addressed them, and the lessons I took away."
)

# Project overview
add_heading("Project Overview", level=2)
add_paragraph(
    "The `Distance` project aims to compute distances and present a small dashboard for geospatial lookups. Key modules in the repository include `dashboard.py`, `geocache.py`, and `main.py`. The project required integrating geocoding, caching results, computing distances between points, and building a user-facing dashboard."
)

# Challenges
add_heading("Challenges Faced", level=2)
add_paragraph(
    "1. Understanding geocoding APIs and rate limits:\n"
    "I needed to research geocoding providers, handle rate limits, and implement retry and caching strategies to avoid excessive API calls. This required reading docs and experimenting with requests.")
add_paragraph(
    "2. Implementing an efficient cache (in `geocache.py`):\n"
    "I designed a local caching layer to persist geocoding results and reduce redundant requests. Debugging cache invalidation and serialization took time to get right.")
add_paragraph(
    "3. Data inconsistencies and edge cases:\n"
    "Real-world location data contained unexpected formats, misspellings, and incomplete addresses. Handling these cases required defensive parsing, normalization, and generous logging.")
add_paragraph(
    "4. Dashboard integration and UX considerations (in `dashboard.py`):\n"
    "Designing a minimal dashboard that is both informative and responsive forced me to iterate on layout, error handling, and clear user feedback for long-running geocoding tasks.")
add_paragraph(
    "5. Time management and scope control:\n"
    "Balancing feature ambition with deadlines meant prioritizing core functionality first (accurate distance calculations and reliable caching) and deferring polish tasks.")

# Solutions & How I fixed things
add_heading("Solutions and Iterations", level=2)
add_paragraph(
    "To address the issues above I took the following steps:\n"
    "• Implemented a robust caching mechanism to store API results and avoid repeated requests.\n"
    "• Added input validation and normalization for addresses to reduce geocoding errors.\n"
    "• Improved logging across modules to make debugging easier and to trace failing inputs.\n"
    "• Broke the dashboard work into small, testable tasks and focused on delivering the critical path first.\n"
)

# Learnings
add_heading("Key Learnings", level=2)
add_paragraph(
    "This project taught me the importance of designing for real data: handling messy inputs, implementing caching early, and providing clear feedback to users. Technically, I became more comfortable with geospatial libraries and HTTP-based APIs. I also improved my debugging workflow and gained experience prioritizing features under time constraints."
)

# Conclusion
add_heading("Conclusion", level=2)
add_paragraph(
    "Overall, the `Distance` project was a valuable learning experience. The challenges I faced helped me grow as a developer and taught me practical strategies for building resilient data-driven tools. I am proud of the progress made and excited about future improvements such as expanding test coverage, refining the UI, and supporting additional data sources."
)

# Footer
add_paragraph("\nSubmitted with the Distance project.\n")

# Save
doc.save(filename)
print(f"Created {filename}")
